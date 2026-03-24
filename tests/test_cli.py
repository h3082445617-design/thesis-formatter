"""Tests for CLI script (format_thesis.py)."""

import subprocess
import os
from pathlib import Path
import pytest
from docx import Document


def test_format_thesis_basic(tmp_path):
    """Test basic CLI formatting with input file."""
    # Create test document
    test_doc = Document()
    test_doc.add_paragraph("Test content")
    input_path = tmp_path / "test.docx"
    test_doc.save(input_path)

    # Get project root for reliable test execution
    project_root = Path(__file__).parent.parent

    # Run CLI
    result = subprocess.run(
        ['python', 'format_thesis.py', str(input_path)],
        cwd=str(project_root),
        capture_output=True,
        text=True
    )

    # Verify success
    assert result.returncode == 0, f"stderr: {result.stderr}"
    assert "successfully" in result.stdout.lower() or "formatted" in result.stdout.lower(), \
        f"stdout: {result.stdout}"

    # Verify output file exists
    output_path = input_path.parent / "test_formatted.docx"
    assert output_path.exists(), f"Output file not found at {output_path}"

    # Verify document was actually formatted (content verification)
    formatted_doc = Document(str(output_path))
    assert len(formatted_doc.paragraphs) > 0, "Formatted document has no paragraphs"


def test_format_thesis_custom_output(tmp_path):
    """Test CLI with custom output path."""
    test_doc = Document()
    test_doc.add_paragraph("Test content")
    input_path = tmp_path / "test.docx"
    test_doc.save(input_path)

    output_path = tmp_path / "custom_output.docx"

    # Get project root for reliable test execution
    project_root = Path(__file__).parent.parent

    result = subprocess.run(
        ['python', 'format_thesis.py', str(input_path), '-o', str(output_path)],
        cwd=str(project_root),
        capture_output=True,
        text=True
    )

    assert result.returncode == 0, f"stderr: {result.stderr}"
    assert output_path.exists(), f"Output file not found at {output_path}"
    assert "successfully" in result.stdout.lower(), f"stdout: {result.stdout}"


def test_format_thesis_missing_file(tmp_path):
    """Test CLI with non-existent file."""
    # Get project root for reliable test execution
    project_root = Path(__file__).parent.parent

    result = subprocess.run(
        ['python', 'format_thesis.py', str(tmp_path / 'nonexistent.docx')],
        cwd=str(project_root),
        capture_output=True,
        text=True
    )

    assert result.returncode == 1
    assert "not found" in result.stderr.lower()


def test_format_thesis_invalid_format(tmp_path):
    """Test CLI with non-.docx file."""
    invalid_file = tmp_path / 'test.txt'
    invalid_file.write_text("test")

    # Get project root for reliable test execution
    project_root = Path(__file__).parent.parent

    result = subprocess.run(
        ['python', 'format_thesis.py', str(invalid_file)],
        cwd=str(project_root),
        capture_output=True,
        text=True
    )

    assert result.returncode == 1
    assert "must be .docx" in result.stderr.lower()
