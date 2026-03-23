import pytest
import os
from docx import Document
from docx.shared import Pt, RGBColor
from src.formatter import DocumentFormatter
from src.exceptions import FormattingError


class TestDocumentFormatter:

    @pytest.fixture
    def sample_docx(self, tmp_path):
        """Create a simple test Word document."""
        doc = Document()
        para = doc.add_paragraph("正文内容测试")
        para.runs[0].font.name = 'SimHei'
        para.runs[0].font.size = Pt(14)

        doc.save(str(tmp_path / "test.docx"))
        return str(tmp_path / "test.docx")

    def test_unify_font_and_line_spacing(self, sample_docx):
        """Test that font is unified to Song (宋体) 12pt and line spacing to 1.5."""
        formatter = DocumentFormatter()
        result_path = formatter.format_document(sample_docx, keep_original=False)

        doc = Document(result_path)

        # Check first paragraph
        para = doc.paragraphs[0]
        for run in para.runs:
            assert run.font.name in ['宋体', 'SimSun', None]  # Chinese fonts may use different names
            assert run.font.size == Pt(12)

        # Check line spacing
        assert para.paragraph_format.line_spacing == 1.5

        # Cleanup
        os.remove(result_path)

    def test_invalid_file_format(self, tmp_path):
        """Test that non-.docx files raise error."""
        txt_file = str(tmp_path / "test.txt")
        with open(txt_file, 'w') as f:
            f.write("test")

        formatter = DocumentFormatter()

        with pytest.raises(FormattingError):
            formatter.format_document(txt_file)

        os.remove(txt_file)
