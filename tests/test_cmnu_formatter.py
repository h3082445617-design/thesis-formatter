# tests/test_cmnu_formatter.py
"""CMNU 格式化器测试"""

import pytest
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from src.formatters import CmnuFormatter
from src.utils import SectionDetector, StyleApplier


class TestCmnuFormatter:
    """CMNU 格式化器集成测试"""

    @pytest.fixture
    def sample_doc(self):
        """创建样本论文文档"""
        doc = Document()
        # TODO: 构建包含各部分的样本文档
        return doc

    def test_formatter_initialization(self):
        """测试格式化器初始化"""
        formatter = CmnuFormatter()

        assert formatter is not None
        assert hasattr(formatter, 'format_document')
        assert hasattr(formatter, '_format_abstract')
        assert hasattr(formatter, '_format_body')

    def test_format_simple_document(self, tmp_path):
        """测试格式化简单文档"""
        # 创建样本文档
        input_doc = Document()
        input_doc.add_paragraph('摘要')
        input_doc.add_paragraph('这是摘要内容')
        input_doc.add_paragraph('关键词: 关键词1; 关键词2')
        input_doc.add_paragraph('第1章 绪论')
        input_doc.add_paragraph('正文内容')

        # 保存为临时文件
        input_path = tmp_path / "test_input.docx"
        input_doc.save(str(input_path))

        # 格式化
        formatter = CmnuFormatter()
        output_path = formatter.format_document(str(input_path))

        # 验证输出文件存在
        assert Path(output_path).exists()

        # 验证输出文档可以打开
        output_doc = Document(output_path)
        assert len(output_doc.paragraphs) > 0

    def test_format_nonexistent_file(self):
        """测试处理不存在的文件"""
        formatter = CmnuFormatter()
        with pytest.raises(FileNotFoundError):
            formatter.format_document("/nonexistent/path/file.docx")

    def test_format_invalid_input_path(self):
        """测试无效的输入路径"""
        formatter = CmnuFormatter()
        with pytest.raises(ValueError):
            formatter.format_document(None)

        with pytest.raises(ValueError):
            formatter.format_document("")

    def test_format_abstract_title(self, tmp_path):
        """测试摘要标题格式化（三号黑体，居中）"""
        doc = Document()
        para_title = doc.add_paragraph('摘要')
        para_content = doc.add_paragraph('这是摘要内容。')
        para_keywords = doc.add_paragraph('关键词: 词1; 词2')

        input_path = tmp_path / "test_abstract.docx"
        doc.save(str(input_path))

        formatter = CmnuFormatter()
        parts = {'abstract': (0, 2)}

        # 手动调用格式化（暂不通过主流程）
        doc = Document(str(input_path))
        formatter._format_abstract(doc, parts)

        # 检查标题
        title_para = doc.paragraphs[0]
        title_run = title_para.runs[0]

        assert title_run.font.name == '黑体'
        assert title_run.font.size == Pt(16)  # 三号
        assert title_run.bold == True
        assert title_para.paragraph_format.alignment == 1  # 居中

    def test_format_abstract_content(self, tmp_path):
        """测试摘要正文格式化（小四号宋体，1.5倍行距，首行缩进）"""
        doc = Document()
        doc.add_paragraph('摘要')
        para_content = doc.add_paragraph('这是摘要内容。')

        input_path = tmp_path / "test_abstract.docx"
        doc.save(str(input_path))

        formatter = CmnuFormatter()
        parts = {'abstract': (0, 1)}

        doc = Document(str(input_path))
        formatter._format_abstract(doc, parts)

        # 检查正文
        content_para = doc.paragraphs[1]
        content_run = content_para.runs[0]

        assert content_run.font.name == '宋体'
        assert content_run.font.size == Pt(12)  # 小四号
        assert content_para.paragraph_format.line_spacing == 1.5
        assert content_para.paragraph_format.first_line_indent == Cm(0.5)  # 2字符缩进

    def test_format_abstract_empty_runs(self, tmp_path):
        """测试空段落不会导致格式化崩溃"""
        doc = Document()
        doc.add_paragraph('摘要')
        # 添加空段落（会有 runs=[] 或类似情况）
        p = doc.add_paragraph()
        p.add_run('')
        doc.add_paragraph('这是摘要内容。')

        input_path = tmp_path / "test_abstract_empty.docx"
        doc.save(str(input_path))

        formatter = CmnuFormatter()
        parts = {'abstract': (0, 2)}

        doc = Document(str(input_path))
        # 不应该崩溃
        formatter._format_abstract(doc, parts)
        assert len(doc.paragraphs) == 3

    def test_format_abstract_missing_part(self, tmp_path):
        """测试摘要部分缺失时的行为"""
        doc = Document()
        doc.add_paragraph('正文内容')

        input_path = tmp_path / "test_no_abstract.docx"
        doc.save(str(input_path))

        formatter = CmnuFormatter()
        parts = {'abstract': None}  # 摘要部分不存在

        doc = Document(str(input_path))
        # 应该提前返回，不出错
        formatter._format_abstract(doc, parts)
        assert len(doc.paragraphs) == 1

    def test_format_abstract_invalid_range(self, tmp_path):
        """测试无效范围的处理"""
        doc = Document()
        doc.add_paragraph('摘要')

        input_path = tmp_path / "test_invalid_range.docx"
        doc.save(str(input_path))

        formatter = CmnuFormatter()
        parts = {'abstract': (100, 200)}  # 超出范围

        doc = Document(str(input_path))
        # 应该优雅地处理，不出错
        formatter._format_abstract(doc, parts)
        # 如果没有错误，测试通过

    # 新增：正文部分的标题检测测试
    def test_detect_heading_level_1(self):
        """测试一级标题检测: '第1章 绪论'"""
        formatter = CmnuFormatter()
        assert formatter._detect_heading_level('第1章 绪论') == 1
        assert formatter._detect_heading_level('第2章 文献综述') == 1
        assert formatter._detect_heading_level('1 绪论') == 1
        assert formatter._detect_heading_level('2 文献综述') == 1

    def test_detect_heading_level_2(self):
        """测试二级标题检测: '1.1 研究背景'"""
        formatter = CmnuFormatter()
        assert formatter._detect_heading_level('1.1 研究背景') == 2
        assert formatter._detect_heading_level('1.1. 研究背景') == 2
        assert formatter._detect_heading_level('2.3 理论基础') == 2

    def test_detect_heading_level_3(self):
        """测试三级标题检测: '1.1.1 问题阐述'"""
        formatter = CmnuFormatter()
        assert formatter._detect_heading_level('1.1.1 问题阐述') == 3
        assert formatter._detect_heading_level('1.1.1. 问题阐述') == 3
        assert formatter._detect_heading_level('2.3.5 细节分析') == 3

    def test_detect_heading_level_0(self):
        """测试非标题文本: '这是普通正文文本。'"""
        formatter = CmnuFormatter()
        assert formatter._detect_heading_level('这是普通正文文本。') == 0
        assert formatter._detect_heading_level('在这个研究中，我们发现了新的现象。') == 0
        assert formatter._detect_heading_level('') == 0

    # 新增：参考文献格式化测试
    def test_format_references_title(self, tmp_path):
        """测试参考文献标题格式化（三号黑体，加粗，居中，1.0倍行距）"""
        doc = Document()
        para_title = doc.add_paragraph('参考文献')
        para_ref1 = doc.add_paragraph('[1] 作者1. 论文题目. 出版社, 2020.')
        para_ref2 = doc.add_paragraph('[2] 作者2. 另一篇论文. 期刊名, 2021.')

        input_path = tmp_path / "test_references.docx"
        doc.save(str(input_path))

        formatter = CmnuFormatter()
        parts = {'references': (0, 2)}

        doc = Document(str(input_path))
        formatter._format_references(doc, parts)

        # 检查标题格式
        title_para = doc.paragraphs[0]
        title_run = title_para.runs[0]

        assert title_run.font.name == '黑体'
        assert title_run.font.size == Pt(16)  # 三号
        assert title_run.bold == True
        assert title_para.paragraph_format.alignment == 1  # 居中
        assert title_para.paragraph_format.line_spacing == 1.0

    def test_format_references_entry(self, tmp_path):
        """测试参考文献条目格式化（五号宋体，1.5倍行距，悬挂缩进，两端对齐）"""
        doc = Document()
        para_title = doc.add_paragraph('参考文献')
        para_ref1 = doc.add_paragraph('[1] 作者1. 论文题目. 出版社, 2020.')
        para_ref2 = doc.add_paragraph('[2] 作者2. 另一篇论文. 期刊名, 2021.')

        input_path = tmp_path / "test_references_entry.docx"
        doc.save(str(input_path))

        formatter = CmnuFormatter()
        parts = {'references': (0, 2)}

        doc = Document(str(input_path))
        formatter._format_references(doc, parts)

        # 检查第一条参考文献
        ref_para = doc.paragraphs[1]
        ref_run = ref_para.runs[0]

        assert ref_run.font.name == '宋体'
        assert ref_run.font.size == Pt(10.5)  # 五号
        assert ref_para.paragraph_format.line_spacing == 1.5
        assert ref_para.paragraph_format.alignment == 3  # 两端对齐

        # 检查悬挂缩进
        pf = ref_para.paragraph_format
        # 悬挂缩进：first_line_indent 为负值，left_indent 为正值
        assert pf.first_line_indent == -Cm(0.64)
        assert pf.left_indent == Cm(0.64)

    def test_format_references_missing_part(self, tmp_path):
        """测试参考文献部分缺失时的行为"""
        doc = Document()
        doc.add_paragraph('正文内容')

        input_path = tmp_path / "test_no_references.docx"
        doc.save(str(input_path))

        formatter = CmnuFormatter()
        parts = {'references': None}  # 参考文献部分不存在

        doc = Document(str(input_path))
        # 应该提前返回，不出错
        formatter._format_references(doc, parts)
        assert len(doc.paragraphs) == 1

    # 新增：正文格式化测试
    def test_format_body_heading_level_1(self, tmp_path):
        """测试一级标题格式化（黑体Pt(16)，加粗，居中）"""
        doc = Document()
        para = doc.add_paragraph('第1章 绪论')

        input_path = tmp_path / "test_body_h1.docx"
        doc.save(str(input_path))

        formatter = CmnuFormatter()
        parts = {'body': (0, 0)}

        doc = Document(str(input_path))
        formatter._format_body(doc, parts)

        # 检查标题格式
        title_para = doc.paragraphs[0]
        assert title_para.paragraph_format.alignment == 1  # 居中
        if title_para.runs:
            title_run = title_para.runs[0]
            assert title_run.font.name == '黑体'
            assert title_run.font.size == Pt(16)  # 三号
            assert title_run.bold == True

    def test_format_body_heading_level_2(self, tmp_path):
        """测试二级标题格式化（黑体Pt(14)，加粗，左对齐）"""
        doc = Document()
        para = doc.add_paragraph('1.1 研究背景')

        input_path = tmp_path / "test_body_h2.docx"
        doc.save(str(input_path))

        formatter = CmnuFormatter()
        parts = {'body': (0, 0)}

        doc = Document(str(input_path))
        formatter._format_body(doc, parts)

        # 检查标题格式
        title_para = doc.paragraphs[0]
        assert title_para.paragraph_format.alignment == 0  # 左对齐
        if title_para.runs:
            title_run = title_para.runs[0]
            assert title_run.font.name == '黑体'
            assert title_run.font.size == Pt(14)  # 四号
            assert title_run.bold == True

    def test_page_numbers_roman_numerals(self, tmp_path):
        """测试页码设置：Roman numerals in front matter"""
        doc = Document()
        # 前文：摘要
        doc.add_paragraph('摘要')
        doc.add_paragraph('这是摘要内容')
        doc.add_paragraph('关键词: 关键词')
        # 正文：从第4段开始
        doc.add_paragraph('第1章 绪论')
        doc.add_paragraph('正文内容')

        input_path = tmp_path / "test_page_numbers.docx"
        doc.save(str(input_path))

        # 打开并设置页码
        doc = Document(str(input_path))
        formatter = CmnuFormatter()
        parts = {'abstract': (0, 2), 'body': (3, 4)}

        # 调用设置页码函数
        formatter._setup_page_numbers(doc, parts)

        # 验证文档包含多个部分（用于不同的页码格式）
        # 前文应该使用罗马数字，正文应该使用阿拉伯数字
        assert len(doc.sections) >= 1

        # 验证Roman format was applied to first section
        section0 = doc.sections[0]
        pgNumType = section0._sectPr.find(qn('w:pgNumType'))
        assert pgNumType is not None, "First section should have pgNumType element"
        fmt_value = pgNumType.get(qn('w:fmt'))
        assert fmt_value == 'upperRoman', f"Expected 'upperRoman', got {fmt_value}"

        # 保存以验证设置
        doc.save(str(tmp_path / "test_page_numbers_output.docx"))

    def test_page_numbers_arabic_in_body(self, tmp_path):
        """测试页码设置：Arabic numerals in body, restarting at 1"""
        doc = Document()
        # 前文：3段
        doc.add_paragraph('摘要')
        doc.add_paragraph('摘要内容第一段')
        doc.add_paragraph('摘要内容第二段')
        # 正文：从第4段开始
        doc.add_paragraph('第1章 绪论')
        doc.add_paragraph('第1章第1节')
        doc.add_paragraph('正文内容')

        input_path = tmp_path / "test_page_numbers_arabic.docx"
        doc.save(str(input_path))

        # 打开并设置页码
        doc = Document(str(input_path))
        formatter = CmnuFormatter()
        parts = {'abstract': (0, 2), 'body': (3, 5)}

        # 调用设置页码函数
        formatter._setup_page_numbers(doc, parts)

        # 验证文档可以保存
        output_path = str(tmp_path / "test_page_numbers_arabic_output.docx")
        doc.save(output_path)

        # 验证输出文件有效
        output_doc = Document(output_path)
        assert len(output_doc.paragraphs) == 6

        # Verify section[0] has Roman format
        section0 = output_doc.sections[0]
        pgNumType0 = section0._sectPr.find(qn('w:pgNumType'))
        assert pgNumType0 is not None, "First section should have pgNumType element"
        fmt_value0 = pgNumType0.get(qn('w:fmt'))
        assert fmt_value0 == 'upperRoman', f"Section[0] expected 'upperRoman', got {fmt_value0}"

        # Verify section[1] has Arabic format with restart at 1
        if len(output_doc.sections) > 1:
            section1 = output_doc.sections[1]
            pgNumType1 = section1._sectPr.find(qn('w:pgNumType'))
            assert pgNumType1 is not None, "Second section should have pgNumType element"
            fmt_value1 = pgNumType1.get(qn('w:fmt'))
            assert fmt_value1 == 'decimal', f"Section[1] expected 'decimal', got {fmt_value1}"
            start_value = pgNumType1.get(qn('w:start'))
            assert start_value == '1', f"Section[1] expected start='1', got {start_value}"

    def test_page_numbers_no_body_part(self, tmp_path):
        """测试页码设置：body part not found, should return early"""
        doc = Document()
        doc.add_paragraph('摘要')
        doc.add_paragraph('摘要内容')

        input_path = tmp_path / "test_no_body.docx"
        doc.save(str(input_path))

        doc = Document(str(input_path))
        formatter = CmnuFormatter()
        parts = {'abstract': (0, 1)}  # 没有 body 部分

        # 应该提前返回，不出错
        formatter._setup_page_numbers(doc, parts)

        # 文档应该仍然有效
        assert len(doc.paragraphs) == 2

        # Verify no new section was created (should still have only 1 section)
        assert len(doc.sections) == 1, "Should still have only 1 section when body part is missing"


class TestSectionDetector:
    """部分检测器测试"""

    def test_match_keywords_abstract(self):
        """测试摘要关键词匹配"""
        doc = Document()
        doc.add_paragraph('摘要')
        doc.add_paragraph('这是摘要内容')

        detector = SectionDetector(doc)
        idx = detector._match_keywords('abstract')

        assert idx == 0, "应该匹配第一段落的'摘要'"

    def test_match_keywords_references(self):
        """测试参考文献关键词匹配"""
        doc = Document()
        doc.add_paragraph('正文内容')
        doc.add_paragraph('参考文献')
        doc.add_paragraph('[1] 作者. 标题')

        detector = SectionDetector(doc)
        idx = detector._match_keywords('references')

        assert idx == 1, "应该匹配第二段落的'参考文献'"

    def test_match_keywords_not_found(self):
        """测试未找到关键词"""
        doc = Document()
        doc.add_paragraph('这是一个文档')
        doc.add_paragraph('没有任何特殊关键词')

        detector = SectionDetector(doc)
        idx = detector._match_keywords('abstract')

        assert idx is None, "未找到时应返回 None"

    def test_detect_parts_simple(self):
        """测试简单场景：检测摘要、目录、正文、参考文献"""
        doc = Document()
        doc.add_paragraph('摘要')
        doc.add_paragraph('摘要正文第一段')
        doc.add_paragraph('摘要正文第二段')
        doc.add_paragraph('关键词: 生物多样性')
        doc.add_paragraph('目录')
        doc.add_paragraph('第1章 ... 1')
        doc.add_paragraph('第1章 正文开始')
        doc.add_paragraph('第1章 正文内容')
        doc.add_paragraph('参考文献')
        doc.add_paragraph('[1] 作者. 标题. 期刊')
        doc.add_paragraph('致谢')
        doc.add_paragraph('感谢导师的指导')

        detector = SectionDetector(doc)
        parts = detector.detect_parts()

        assert parts['abstract'] == (0, 3), "摘要应该是 0-3"
        assert parts['toc'] == (4, 5), "目录应该是 4-5"
        assert parts['body'] == (6, 8), "正文应该是 6-8"
        assert parts['references'] == (9, 10), "参考文献应该是 9-10"
        assert parts['acknowledgment'] == (11, 12), "致谢应该是 11-12"

    def test_detect_parts_missing_toc(self):
        """测试目录缺失的情况"""
        doc = Document()
        doc.add_paragraph('摘要')
        doc.add_paragraph('摘要内容')
        doc.add_paragraph('第1章 标题')  # 直接开始正文，没有目录
        doc.add_paragraph('正文内容')
        doc.add_paragraph('参考文献')

        detector = SectionDetector(doc)
        parts = detector.detect_parts()

        # 应该至少检测到摘要、正文、参考文献
        assert parts['abstract'] is not None
        assert parts['body'] is not None
        assert parts['references'] is not None
        assert parts['toc'] is None  # 目录未检测到

    def test_detect_abstract(self):
        """测试摘要检测"""
        # TODO: 实现
        pass

    def test_detect_references(self):
        """测试参考文献检测"""
        # TODO: 实现
        pass


class TestStyleApplier:
    """样式应用器测试"""

    def test_apply_font_song(self):
        """测试应用宋体"""
        doc = Document()
        para = doc.add_paragraph('测试文本')
        run = para.runs[0]

        applier = StyleApplier()
        applier.apply_font(run, '宋体', Pt(12), bold=False)

        assert run.font.name == '宋体'
        assert run.font.size == Pt(12)
        assert run.bold == False

    def test_apply_font_hei_bold(self):
        """测试应用黑体加粗"""
        doc = Document()
        para = doc.add_paragraph('标题')
        run = para.runs[0]

        applier = StyleApplier()
        applier.apply_font(run, '黑体', Pt(16), bold=True)

        assert run.font.name == '黑体'
        assert run.font.size == Pt(16)
        assert run.bold == True

    def test_apply_paragraph_format_standard(self):
        """测试应用标准段落格式（1.5倍行距、首行缩进）"""
        doc = Document()
        para = doc.add_paragraph('正文内容')

        applier = StyleApplier()
        applier.apply_paragraph_format(
            para,
            alignment=0,  # 左对齐
            line_spacing=1.5,
            space_before=Pt(0),
            space_after=Pt(6),
            indent_first_line=Cm(0.5)
        )

        pf = para.paragraph_format
        assert pf.line_spacing == 1.5
        assert pf.space_before == Pt(0)
        assert pf.space_after == Pt(6)
        assert pf.first_line_indent == Cm(0.5)

    def test_apply_hanging_indent(self):
        """测试参考文献悬挂缩进"""
        doc = Document()
        para = doc.add_paragraph('[1] 作者. 标题. 期刊')

        applier = StyleApplier()
        applier.apply_hanging_indent(para, hanging_indent=Cm(0.5))

        pf = para.paragraph_format
        # 悬挂缩进: first_line_indent = -hanging_indent
        assert pf.first_line_indent == Cm(-0.5)
        assert pf.left_indent == Cm(0.5)
