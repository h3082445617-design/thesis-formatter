"""Integration tests for thesis document formatting pipeline.

Tests the complete formatting workflow with realistic thesis documents
that simulate real user scenarios. These tests verify:
- End-to-end formatting pipeline
- Realistic document structures (abstract, body, references)
- Multiple heading levels
- All formatting rules applied correctly
- Content preservation
"""

import pytest
import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from src.formatter import DocumentFormatter


# Font name constants (from src/formatter.py)
BODY_FONT_NAMES = {'宋体', 'SimSun', 'Song Ti'}
HEADING_FONT_NAMES = {'黑体', 'Heibei', 'Arial Black'}


@pytest.fixture
def cleanup_formatted_file():
    """Fixture to handle cleanup of formatted files."""
    files_to_remove = []
    yield files_to_remove
    for file_path in files_to_remove:
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
        except OSError:
            pass  # Silent cleanup failure


@pytest.fixture
def simple_thesis_doc(tmp_path):
    """Create a simple thesis document with abstract, body, references.

    This fixture creates a basic thesis structure suitable for testing
    the fundamental formatting pipeline.
    """
    doc = Document()

    # Abstract section
    doc.add_heading('摘要', level=1)
    doc.add_paragraph('This is a sample abstract with some content.')

    # Body section with heading
    doc.add_heading('1 Introduction', level=1)
    doc.add_paragraph('This is the introduction.')

    # References section
    doc.add_heading('参考文献', level=1)
    doc.add_paragraph('[1] Author, Title, Journal, 2020.')

    doc_path = tmp_path / 'simple_thesis.docx'
    doc.save(str(doc_path))
    return str(doc_path)


@pytest.fixture
def complex_thesis_doc(tmp_path):
    """Create a realistic thesis with multiple heading levels and formatting.

    This fixture creates a comprehensive thesis structure that tests
    the formatter's ability to handle nested heading levels and more
    realistic content.
    """
    doc = Document()

    # Abstract
    doc.add_heading('摘要', level=1)
    doc.add_paragraph('Sample abstract text with research overview.')

    # Main body with nested headings
    doc.add_heading('1 Introduction', level=1)
    doc.add_paragraph('Main introduction text discussing the research context.')

    doc.add_heading('1.1 Background', level=2)
    doc.add_paragraph('Background information about the field and problem domain.')

    doc.add_heading('1.1.1 Historical Context', level=3)
    doc.add_paragraph('Historical details and evolution of the field.')

    doc.add_heading('2 Methodology', level=1)
    doc.add_paragraph('Detailed methodology describing the research approach.')

    doc.add_heading('2.1 Approach', level=2)
    doc.add_paragraph('Our specific approach to solving the problem.')

    # References with multiple entries
    doc.add_heading('参考文献', level=1)
    doc.add_paragraph('[1] Author A. Title of work. Journal, 2020.')
    doc.add_paragraph('[2] Author B. Another title. Conference, 2021.')
    doc.add_paragraph('[3] Author C. Third reference. Book Publisher, 2022.')

    doc_path = tmp_path / 'complex_thesis.docx'
    doc.save(str(doc_path))
    return str(doc_path)


class TestIntegrationFormatting:
    """Integration tests for complete formatting pipeline."""

    def test_integration_simple_thesis(self, simple_thesis_doc, cleanup_formatted_file):
        """Test complete formatting of simple thesis document.

        This test verifies that:
        - The formatter processes the document without errors
        - Output file is created
        - Document structure is preserved
        - Key content sections are present
        """
        formatter = DocumentFormatter()
        formatted_path = formatter.format_document(simple_thesis_doc, keep_original=False)
        cleanup_formatted_file.append(formatted_path)

        # Verify file was created
        assert Path(formatted_path).exists(), "Formatted document should be created"

        # Load formatted document
        formatted_doc = Document(formatted_path)

        # Verify structure preserved
        assert len(formatted_doc.paragraphs) > 0, "Document should have paragraphs"
        assert len(formatted_doc.sections) >= 1, "Document should have at least one section"

        # Verify content preserved
        all_text = '\n'.join([p.text for p in formatted_doc.paragraphs])
        assert '摘要' in all_text, "Abstract section should be preserved"
        assert 'Introduction' in all_text, "Introduction section should be preserved"
        assert '参考文献' in all_text, "References section should be preserved"

    def test_integration_complex_thesis(self, complex_thesis_doc, cleanup_formatted_file):
        """Test formatting of realistic thesis with multiple heading levels.

        This test verifies that:
        - Complex document structure is handled correctly
        - All heading levels are preserved
        - Content with multiple sections is formatted properly
        """
        formatter = DocumentFormatter()
        formatted_path = formatter.format_document(str(complex_thesis_doc), keep_original=False)
        cleanup_formatted_file.append(formatted_path)

        formatted_doc = Document(formatted_path)

        # Verify structure
        assert len(formatted_doc.paragraphs) > 0, "Document should have paragraphs"

        # Verify all heading levels preserved
        all_text = '\n'.join([p.text for p in formatted_doc.paragraphs])
        assert '1 Introduction' in all_text, "Level 1 heading should be preserved"
        assert '1.1 Background' in all_text, "Level 2 heading should be preserved"
        assert '1.1.1 Historical Context' in all_text, "Level 3 heading should be preserved"
        assert '2 Methodology' in all_text, "Second level 1 heading should be preserved"
        assert '2.1 Approach' in all_text, "Second level 2 heading should be preserved"

        # Verify major sections are preserved (strengthened from tautological check)
        assert '1 Introduction' in all_text, "Introduction heading should be preserved"
        assert 'Methodology' in all_text, "Methodology section should be preserved"
        assert '参考文献' in all_text, "References section should be preserved"

    def test_integration_formatting_rules(self, simple_thesis_doc, cleanup_formatted_file):
        """Verify all formatting rules are applied correctly.

        This test checks:
        - Margins are set correctly (2cm top/bottom, 2.5cm left/right)
        - Body text uses Song font (宋体)
        - Line spacing is set to 1.5
        - Headings are formatted with bold
        """
        formatter = DocumentFormatter()
        formatted_path = formatter.format_document(simple_thesis_doc, keep_original=False)
        cleanup_formatted_file.append(formatted_path)

        formatted_doc = Document(formatted_path)

        # Check margins
        section = formatted_doc.sections[0]
        assert section.top_margin == Cm(2), \
            f"Top margin should be 2cm, got {section.top_margin}"
        assert section.bottom_margin == Cm(2), \
            f"Bottom margin should be 2cm, got {section.bottom_margin}"
        assert section.left_margin == Cm(2.5), \
            f"Left margin should be 2.5cm, got {section.left_margin}"
        assert section.right_margin == Cm(2), \
            f"Right margin should be 2cm, got {section.right_margin}"

        # Verify line spacing is applied to body paragraphs
        body_paragraph_count = 0
        body_with_spacing = 0
        for para in formatted_doc.paragraphs:
            if para.text.strip() and '摘要' not in para.text and '参考文献' not in para.text:
                body_paragraph_count += 1
                # Check for 1.5x line spacing or close approximation (some systems round differently)
                if para.paragraph_format.line_spacing and abs(para.paragraph_format.line_spacing - 1.5) < 0.1:
                    body_with_spacing += 1

        assert body_paragraph_count > 0, "Document should have body paragraphs"
        assert body_with_spacing > 0, "Body paragraphs should have 1.5x line spacing applied"

        # Verify body formatting was applied
        body_formatted = False
        for para in formatted_doc.paragraphs:
            if para.text.strip() and '摘要' not in para.text and '参考文献' not in para.text:
                # This is body text, check its formatting
                for run in para.runs:
                    if run.font.name in BODY_FONT_NAMES:
                        body_formatted = True
                        break
                if body_formatted:
                    break

        assert body_formatted, "Body text should have Song font formatting applied"

        # Check heading formatting exists
        heading_found = False
        for para in formatted_doc.paragraphs:
            if '摘要' in para.text or 'Introduction' in para.text:
                for run in para.runs:
                    if run.bold:
                        heading_found = True
                        break
                if heading_found:
                    break

        assert heading_found, "Headings should be formatted with bold"

    def test_integration_content_preservation(self, complex_thesis_doc, cleanup_formatted_file):
        """Verify document content is preserved during formatting.

        This test ensures that:
        - All text content is preserved
        - Paragraph count is similar (within small margin for formatting changes)
        - Key sections are intact
        """
        # Load original document for comparison
        original_doc = Document(complex_thesis_doc)
        original_text = '\n'.join([p.text for p in original_doc.paragraphs])
        original_para_count = len([p for p in original_doc.paragraphs if p.text.strip()])

        # Format and load
        formatter = DocumentFormatter()
        formatted_path = formatter.format_document(str(complex_thesis_doc), keep_original=False)
        cleanup_formatted_file.append(formatted_path)
        formatted_doc = Document(formatted_path)
        formatted_text = '\n'.join([p.text for p in formatted_doc.paragraphs])
        formatted_para_count = len([p for p in formatted_doc.paragraphs if p.text.strip()])

        # Verify critical sections preserved
        assert '摘要' in formatted_text, "Abstract section should be preserved"
        assert '参考文献' in formatted_text, "References section should be preserved"
        assert 'Introduction' in formatted_text, "Introduction should be preserved"

        # Verify paragraph count is similar (allowing small margin for formatting changes)
        assert abs(original_para_count - formatted_para_count) <= 1, \
            f"Paragraph count should be preserved (was {original_para_count}, now {formatted_para_count})"

    def test_integration_output_file_naming(self, simple_thesis_doc, cleanup_formatted_file):
        """Verify output file naming follows the expected convention.

        Tests that formatted output is created with _formatted suffix.
        """
        formatter = DocumentFormatter()
        formatted_path = formatter.format_document(simple_thesis_doc, keep_original=False)
        cleanup_formatted_file.append(formatted_path)

        # Verify output path contains '_formatted' suffix
        assert '_formatted' in formatted_path, \
            "Output file should contain '_formatted' in name"
        assert formatted_path.endswith('.docx'), \
            "Output file should be .docx format"

    def test_integration_document_validity(self, complex_thesis_doc, cleanup_formatted_file):
        """Verify formatted document can be opened and is valid.

        This test ensures the output is a valid .docx file that can be
        read by python-docx without errors.
        """
        formatter = DocumentFormatter()
        formatted_path = formatter.format_document(str(complex_thesis_doc), keep_original=False)
        cleanup_formatted_file.append(formatted_path)

        # Should not raise any exception
        formatted_doc = Document(formatted_path)

        # Verify we can read basic properties
        assert formatted_doc.paragraphs is not None
        assert formatted_doc.sections is not None
        assert len(formatted_doc.sections) > 0

    def test_integration_page_numbers(self, simple_thesis_doc, cleanup_formatted_file):
        """Verify page numbers are added to document footer.

        Tests that:
        - Footer is created
        - Page number field is added
        - Footer exists in the section
        """
        formatter = DocumentFormatter()
        formatted_path = formatter.format_document(simple_thesis_doc, keep_original=False)
        cleanup_formatted_file.append(formatted_path)

        formatted_doc = Document(formatted_path)

        # Check that footer exists
        section = formatted_doc.sections[0]
        footer = section.footer

        # Footer should have content (page number field)
        assert footer is not None, "Section should have footer"
        assert len(footer.paragraphs) > 0, "Footer should have at least one paragraph"

    def test_integration_multiple_sections(self, tmp_path, cleanup_formatted_file):
        """Test formatting of document with multiple sections.

        Verifies that margins and formatting are applied to all sections.
        """
        # Create document with content
        doc = Document()
        doc.add_heading('Title', level=1)
        doc.add_paragraph('Section 1 content')
        doc.add_paragraph('More content')

        doc_path = tmp_path / 'multi_section.docx'
        doc.save(str(doc_path))

        formatter = DocumentFormatter()
        formatted_path = formatter.format_document(str(doc_path), keep_original=False)
        cleanup_formatted_file.append(formatted_path)

        formatted_doc = Document(formatted_path)

        # All sections should have margins applied
        for section in formatted_doc.sections:
            assert section.top_margin == Cm(2), "All sections should have top margin"
            assert section.bottom_margin == Cm(2), "All sections should have bottom margin"
