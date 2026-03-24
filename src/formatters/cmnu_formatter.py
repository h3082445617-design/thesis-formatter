"""中央民族大学学位论文格式化器"""

import re
from typing import Dict, Tuple, Optional
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.enum.section import WD_SECTION
from src.utils import SectionDetector, StyleApplier


class CmnuFormatter:
    """中央民族大学学位论文格式化器"""

    # CMNU 标准参数
    MARGINS = {
        'top': Cm(2.5),
        'bottom': Cm(2.5),
        'left': Cm(2.2),
        'right': Cm(2.2)
    }

    # 字体大小常量
    FONT_SIZES = {
        '三号': Pt(16),      # 摘要标题
        '四号': Pt(14),      # 关键词标签
        '小四号': Pt(12),    # 关键词内容、正文
        '五号': Pt(10.5),    # 参考文献条目
    }

    def __init__(self):
        """初始化格式化器"""
        self.style_applier = StyleApplier()
        self.debug = False

    def format_document(self, input_path: str, output_path: Optional[str] = None,
                        debug: bool = False) -> str:
        """
        格式化论文文档

        Args:
            input_path: 输入 .docx 文件路径
            output_path: 输出文件路径（默认：input_formatted.docx）
            debug: 是否打印调试信息

        Returns:
            输出文件路径

        Raises:
            ValueError: 如果 input_path 无效
            FileNotFoundError: 如果文件不存在
            IOError: 如果文件操作失败
        """
        self.debug = debug

        # 1. 验证输入路径
        if not input_path or not isinstance(input_path, str):
            raise ValueError(f"input_path must be a non-empty string, got: {input_path}")

        input_file = Path(input_path)
        if not input_file.exists():
            raise FileNotFoundError(f"Input file does not exist: {input_path}")

        # 2. 打开文档
        try:
            doc = Document(input_path)
        except FileNotFoundError:
            raise FileNotFoundError(f"Input file not found: {input_path}")
        except Exception as e:
            raise IOError(f"Failed to open document: {input_path}. Error: {str(e)}")

        # 1. 检测各部分
        detector = SectionDetector(doc)
        parts = detector.detect_parts(debug=self.debug)

        if self.debug:
            print(f"Detected parts: {parts}")

        # 2. 应用边距
        self._apply_margins(doc)

        # 3. 应用各部分格式
        self._format_abstract(doc, parts)
        self._format_toc(doc, parts)
        self._format_body(doc, parts)
        self._format_references(doc, parts)
        self._format_appendix(doc, parts)
        self._format_acknowledgment(doc, parts)

        # 4. 处理页码
        self._setup_page_numbers(doc, parts)

        # 5. 保存输出
        if output_path is None:
            base_path = Path(input_path)
            output_path = str(base_path.parent / f"{base_path.stem}_formatted.docx")

        # 确保输出目录存在
        output_dir = Path(output_path).parent
        output_dir.mkdir(parents=True, exist_ok=True)

        doc.save(output_path)

        if self.debug:
            print(f"Document saved to: {output_path}")

        return output_path

    def _apply_margins(self, doc: Document):
        """应用页边距"""
        for section in doc.sections:
            section.top_margin = self.MARGINS['top']
            section.bottom_margin = self.MARGINS['bottom']
            section.left_margin = self.MARGINS['left']
            section.right_margin = self.MARGINS['right']

    # TODO: 实现各部分格式化方法
    def _format_abstract(self, doc: Document, parts: Dict):
        """
        格式化摘要部分

        标题: 三号黑体，居中，单倍行距，段前24磅/段后18磅
        正文: 小四号宋体，1.5倍行距，首行缩进2字符
        关键词标签: 四号黑体
        关键词内容: 小四号宋体
        """
        # 验证 parts 字典结构
        if parts.get('abstract') is None:
            return

        abstract_part = parts.get('abstract')
        if not isinstance(abstract_part, tuple) or len(abstract_part) != 2:
            if self.debug:
                print(f"Warning: Invalid abstract part definition: {abstract_part}")
            return

        start, end = abstract_part

        for idx in range(start, end + 1):
            if idx >= len(doc.paragraphs):
                break

            para = doc.paragraphs[idx]
            text = para.text.strip()

            # 标题行（包含"摘要"）
            if '摘要' in text and idx == start:
                if para.runs:  # 检查 runs 是否非空
                    for run in para.runs:
                        self.style_applier.apply_font(run, '黑体', self.FONT_SIZES['三号'], bold=True)
                pf = para.paragraph_format
                pf.alignment = 1  # 居中
                pf.line_spacing = 1.0
                pf.space_before = Pt(24)
                pf.space_after = Pt(18)

            # 关键词标签行
            elif text.startswith('关键词:') or text.startswith('关键词：'):
                # 假设：关键词标签和内容在单独的 run 或段落中
                # 处理格式：[keyword-run]"关键词"[content-run]"词1; 词2"
                if para.runs:  # 检查 runs 是否非空
                    for run in para.runs:
                        if '关键词' in run.text:
                            self.style_applier.apply_font(run, '黑体', self.FONT_SIZES['四号'], bold=True)
                        else:
                            self.style_applier.apply_font(run, '宋体', self.FONT_SIZES['小四号'], bold=False)

            # 普通正文
            else:
                if para.runs:  # 检查 runs 是否非空
                    for run in para.runs:
                        self.style_applier.apply_font(run, '宋体', self.FONT_SIZES['小四号'], bold=False)

                pf = para.paragraph_format
                pf.line_spacing = 1.5
                pf.first_line_indent = Cm(0.5)  # 2字符缩进
                pf.space_before = Pt(0)
                pf.space_after = Pt(6)

    def _detect_heading_level(self, text: str) -> int:
        """
        检测段落的标题级别

        Args:
            text: 段落文本

        Returns:
            int: 标题级别
                - 3: 三级标题 (e.g., "1.1.1 ")
                - 2: 二级标题 (e.g., "1.1 ")
                - 1: 一级标题 (e.g., "第1章", "1 ")
                - 0: 普通文本
        """
        if not text or not isinstance(text, str):
            return 0

        text = text.strip()
        if not text:
            return 0

        # 按优先级检测：先检测3级，再检测2级，最后检测1级
        # 三级标题: 1.1.1 或 1.1.1.
        if re.match(r'^\d+\.\d+\.\d+[\.\s]', text):
            return 3

        # 二级标题: 1.1 或 1.1.
        if re.match(r'^\d+\.\d+[\.\s]', text):
            return 2

        # 一级标题: 第1章、第二章 或 1 、2 等
        if re.match(r'^(第[一二三四五六七八九十〇\d]+章|[\d\u4e00-\u9fff]\s)', text):
            return 1

        # 非标题
        return 0

    def _format_toc(self, doc: Document, parts: Dict):
        """格式化目录"""
        pass

    def _format_body(self, doc: Document, parts: Dict):
        """
        格式化正文

        标题格式化规则：
        - 一级标题: 黑体 Pt(16)，加粗，居中，单倍行距，段前24磅/段后18磅
        - 二级标题: 黑体 Pt(14)，加粗，左对齐，20磅行距，段前24磅/段后6磅
        - 三级标题: 黑体 Pt(12)，加粗，左对齐，20磅行距，段前12磅/段后6磅
        - 普通文本: 宋体 Pt(12)，1.5倍行距，0.5cm首行缩进，两端对齐，段前0/段后6磅
        """
        # 验证 parts 字典结构
        if parts.get('body') is None:
            return

        body_part = parts.get('body')
        if not isinstance(body_part, tuple) or len(body_part) != 2:
            if self.debug:
                print(f"Warning: Invalid body part definition: {body_part}")
            return

        start, end = body_part

        for idx in range(start, end + 1):
            if idx >= len(doc.paragraphs):
                break

            para = doc.paragraphs[idx]
            text = para.text.strip()

            # 检测标题级别
            heading_level = self._detect_heading_level(text)

            if heading_level == 1:
                # 一级标题: 黑体 Pt(16)，加粗，居中，单倍行距，段前24磅/段后18磅
                if para.runs:
                    for run in para.runs:
                        self.style_applier.apply_font(run, '黑体', self.FONT_SIZES['三号'], bold=True)
                pf = para.paragraph_format
                pf.alignment = 1  # 居中
                pf.line_spacing = 1.0
                pf.space_before = Pt(24)
                pf.space_after = Pt(18)

            elif heading_level == 2:
                # 二级标题: 黑体 Pt(14)，加粗，左对齐，20磅固定行距，段前24磅/段后6磅
                if para.runs:
                    for run in para.runs:
                        self.style_applier.apply_font(run, '黑体', self.FONT_SIZES['四号'], bold=True)
                pf = para.paragraph_format
                pf.alignment = 0  # 左对齐
                pf.line_spacing = 20  # 20磅固定行距
                pf.space_before = Pt(24)
                pf.space_after = Pt(6)

            elif heading_level == 3:
                # 三级标题: 黑体 Pt(12)，加粗，左对齐，20磅固定行距，段前12磅/段后6磅
                if para.runs:
                    for run in para.runs:
                        self.style_applier.apply_font(run, '黑体', self.FONT_SIZES['小四号'], bold=True)
                pf = para.paragraph_format
                pf.alignment = 0  # 左对齐
                pf.line_spacing = 20  # 20磅固定行距
                pf.space_before = Pt(12)
                pf.space_after = Pt(6)

            else:
                # 普通文本: 宋体 Pt(12)，1.5倍行距，0.5cm首行缩进，两端对齐，段前0/段后6磅
                if para.runs:
                    for run in para.runs:
                        self.style_applier.apply_font(run, '宋体', self.FONT_SIZES['小四号'], bold=False)
                pf = para.paragraph_format
                pf.alignment = 3  # 两端对齐
                pf.line_spacing = 1.5
                pf.first_line_indent = Cm(0.5)
                pf.space_before = Pt(0)
                pf.space_after = Pt(6)

    def _format_references(self, doc: Document, parts: Dict):
        """
        格式化参考文献

        标题: 三号黑体，加粗，居中，单倍行距，段前24磅/段后18磅
        条目: 五号宋体，1.5倍行距，两端对齐，悬挂缩进（-0.25英寸），段前0/段后6磅
        """
        # 验证 parts 字典结构
        if parts.get('references') is None:
            return

        references_part = parts.get('references')
        if not isinstance(references_part, tuple) or len(references_part) != 2:
            if self.debug:
                print(f"Warning: Invalid references part definition: {references_part}")
            return

        start, end = references_part

        for idx in range(start, end + 1):
            if idx >= len(doc.paragraphs):
                break

            para = doc.paragraphs[idx]
            text = para.text.strip()

            # 标题行（参考文献）
            if '参考文献' in text and idx == start:
                if para.runs:  # 检查 runs 是否非空
                    for run in para.runs:
                        self.style_applier.apply_font(run, '黑体', self.FONT_SIZES['三号'], bold=True)
                pf = para.paragraph_format
                pf.alignment = 1  # 居中
                pf.line_spacing = 1.0
                pf.space_before = Pt(24)
                pf.space_after = Pt(18)

            # 参考文献条目
            else:
                if para.runs:  # 检查 runs 是否非空
                    for run in para.runs:
                        self.style_applier.apply_font(run, '宋体', self.FONT_SIZES['五号'], bold=False)

                # 应用段落格式和悬挂缩进
                pf = para.paragraph_format
                pf.alignment = 3  # 两端对齐 (justified)
                pf.line_spacing = 1.5
                pf.space_before = Pt(0)
                pf.space_after = Pt(6)

                # 应用悬挂缩进（-0.25英寸 ≈ -0.64cm）
                self.style_applier.apply_hanging_indent(para, Cm(0.64))

    def _format_appendix(self, doc: Document, parts: Dict):
        """格式化附录"""
        pass

    def _format_acknowledgment(self, doc: Document, parts: Dict):
        """格式化致谢"""
        pass

    def _setup_page_numbers(self, doc: Document, parts: Dict):
        """
        设置双页码系统：前文（罗马数字）和正文（阿拉伯数字）

        Args:
            doc: 文档对象
            parts: 各部分范围字典，如 {'body': (3, 4), ...}
        """
        # 验证 body 部分是否存在
        body_part = parts.get('body')
        if body_part is None:
            if self.debug:
                print("No body part found, skipping page numbering setup")
            return

        if not isinstance(body_part, tuple) or len(body_part) != 2:
            if self.debug:
                print(f"Invalid body part definition: {body_part}")
            return

        body_start, body_end = body_part

        try:
            # 1. 设置第一部分（前文）为罗马数字
            self._setup_section_page_numbers(doc.sections[0], 'roman', None)

            # 2. 在正文开始处创建新部分
            self._create_new_section(doc, body_start)

            # 3. 设置第二部分（正文）为阿拉伯数字，从1开始
            if len(doc.sections) > 1:
                self._setup_section_page_numbers(doc.sections[1], 'arabic', 1)

            if self.debug:
                print(f"Page numbering setup completed: {len(doc.sections)} sections")

        except Exception as e:
            if self.debug:
                print(f"Error setting up page numbers: {str(e)}")
            # 降级方案：应用简单的阿拉伯数字页码到所有部分
            self._setup_simple_page_numbers(doc)

    def _create_new_section(self, doc: Document, at_paragraph_idx: int):
        """
        在指定段落处创建新的部分（section）

        Args:
            doc: 文档对象
            at_paragraph_idx: 创建新部分的段落索引
        """
        if at_paragraph_idx >= len(doc.paragraphs):
            if self.debug:
                print(f"Paragraph index {at_paragraph_idx} out of range")
            return

        paragraph = doc.paragraphs[at_paragraph_idx]
        p = paragraph._element

        # 获取原始 section 属性
        original_sectPr = doc.sections[0]._sectPr

        # 创建新的 section properties
        new_sectPr = OxmlElement('w:sectPr')

        # 复制关键属性：页面尺寸、边距等
        for child in original_sectPr:
            tag_name = child.tag.split('}')[-1] if '}' in child.tag else child.tag
            # 跳过已有的 page number type 设置
            if tag_name not in ['pgNumType', 'footerReference', 'headerReference']:
                new_child = OxmlElement(child.tag)
                new_child.attrib.update(child.attrib)
                new_sectPr.append(new_child)

        # 附加新的 section properties 到段落
        p.append(new_sectPr)

        if self.debug:
            print(f"Created new section at paragraph {at_paragraph_idx}")

    def _setup_section_page_numbers(self, section, num_format: str, start_num: Optional[int] = None):
        """
        为特定部分设置页码格式

        Args:
            section: 文档部分对象
            num_format: 页码格式 ('roman' 或 'arabic')
            start_num: 起始页码（如果为None，继续使用前一部分的编号）
        """
        try:
            sectPr = section._sectPr

            # 创建或修改 page number type 元素
            pgNumType = sectPr.find(qn('w:pgNumType'))
            if pgNumType is None:
                pgNumType = OxmlElement('w:pgNumType')
                sectPr.insert(0, pgNumType)

            # 设置页码格式
            if num_format == 'roman':
                pgNumType.set(qn('w:fmt'), 'upperRoman')
            elif num_format == 'arabic':
                pgNumType.set(qn('w:fmt'), 'decimal')

            # 设置起始页码
            if start_num is not None:
                pgNumType.set(qn('w:start'), str(start_num))

            if self.debug:
                print(f"Section page format set to {num_format}, start={start_num}")

        except Exception as e:
            if self.debug:
                print(f"Error setting up section page numbers: {str(e)}")

    def _setup_simple_page_numbers(self, doc: Document):
        """
        简单的页码设置降级方案：应用阿拉伯数字到所有部分

        Args:
            doc: 文档对象
        """
        try:
            for section in doc.sections:
                self._setup_section_page_numbers(section, 'arabic', None)

            if self.debug:
                print("Applied simple page numbering (Arabic to all sections)")

        except Exception as e:
            if self.debug:
                print(f"Error setting up simple page numbers: {str(e)}")
