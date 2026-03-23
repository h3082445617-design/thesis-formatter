"""Core thesis document formatting engine."""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from src.exceptions import FormattingError, FileCorruptedError, InvalidFileFormatError
import os
import time


class DocumentFormatter:
    """Core engine for formatting thesis documents according to GB/T 7714 standard."""

    TIMEOUT_SECONDS = 30

    def __init__(self):
        """Initialize formatter with GB/T 7714 standard settings."""
        self.body_font_name = '宋体'
        self.body_font_size = Pt(12)
        self.body_line_spacing = 1.5
        self.heading1_font_size = Pt(18)
        self.heading2_font_size = Pt(14)
        self.heading3_font_size = Pt(12)
        self.paragraph_space_before = Pt(0)
        self.paragraph_space_after = Pt(6)
        self.margin_left = Cm(2.5)
        self.margin_right = Cm(2)
        self.margin_top = Cm(2)
        self.margin_bottom = Cm(2)

    def format_document(self, input_path, keep_original=True):
        """
        Format a Word document according to GB/T 7714 standard.

        Args:
            input_path: Path to input .docx file
            keep_original: If True, save original as backup

        Returns:
            Path to formatted document

        Raises:
            InvalidFileFormatError: If file is not .docx
            FileCorruptedError: If file cannot be opened
            FormattingError: If formatting fails
        """
        start_time = time.time()

        # Validate file format
        if not input_path.lower().endswith('.docx'):
            raise InvalidFileFormatError(f"文件格式必须是.docx，你上传的是 {os.path.splitext(input_path)[1]} 文件")

        # Check file exists
        if not os.path.exists(input_path):
            raise FileCorruptedError("文件无法访问或不存在")

        try:
            # Load document
            doc = Document(input_path)
        except Exception as e:
            raise FileCorruptedError(f"文件无效或损坏：{str(e)}")

        try:
            # Apply formatting
            self._unify_fonts_and_spacing(doc)
            self._detect_and_format_titles(doc)
            self._apply_margins(doc)
            self._add_page_numbers(doc)

            # Generate output path
            base, ext = os.path.splitext(input_path)
            output_path = f"{base}_formatted{ext}"

            # Save
            doc.save(output_path)

            elapsed = time.time() - start_time
            if elapsed > self.TIMEOUT_SECONDS:
                raise FormattingError(f"处理超时（{elapsed:.1f}秒），你的文件可能过于复杂")

            return output_path

        except FormattingError:
            raise
        except Exception as e:
            raise FormattingError(f"处理失败：{str(e)}")

    def _unify_fonts_and_spacing(self, doc):
        """Apply uniform font, size, and line spacing to all paragraphs."""
        for paragraph in doc.paragraphs:
            # Set line spacing
            paragraph.paragraph_format.line_spacing = self.body_line_spacing

            # Set paragraph spacing
            paragraph.paragraph_format.space_before = self.paragraph_space_before
            paragraph.paragraph_format.space_after = self.paragraph_space_after

            # Set font for all runs
            for run in paragraph.runs:
                run.font.name = self.body_font_name
                run.font.east_asia_font_name = self.body_font_name  # For Chinese characters
                run.font.size = self.body_font_size

        # Also apply to table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph.paragraph_format.line_spacing = self.body_line_spacing
                        for run in paragraph.runs:
                            run.font.name = self.body_font_name
                            run.font.east_asia_font_name = self.body_font_name
                            run.font.size = self.body_font_size

    def _detect_and_format_titles(self, doc):
        """Detect and format titles by hierarchy (H1, H2, H3)."""
        for i, paragraph in enumerate(doc.paragraphs):
            if not paragraph.runs:
                continue

            # First check Word style
            style_name = paragraph.style.name
            if 'Heading 1' in style_name or style_name == 'Heading 1':
                self._format_heading1(paragraph)
                continue
            elif 'Heading 2' in style_name or style_name == 'Heading 2':
                self._format_heading2(paragraph)
                continue
            elif 'Heading 3' in style_name or style_name == 'Heading 3':
                self._format_heading3(paragraph)
                continue

            # Then try to detect by font size and formatting
            first_run = paragraph.runs[0]
            current_size = first_run.font.size
            is_bold = first_run.bold

            # Detect by font size and bold
            if is_bold or (current_size and current_size >= Pt(16)):
                # Level 1 heading
                self._format_heading1(paragraph)
            elif current_size and current_size >= Pt(13) and is_bold:
                # Level 2 heading
                self._format_heading2(paragraph)
            elif is_bold and current_size and current_size == Pt(12):
                # Level 3 heading (12pt bold)
                self._format_heading3(paragraph)

    def _format_heading1(self, paragraph):
        """Format as level 1 heading: HeiBei 18pt bold centered."""
        for run in paragraph.runs:
            run.font.name = '黑体'
            run.font.east_asia_font_name = '黑体'
            run.font.size = self.heading1_font_size
            run.bold = True
        paragraph.alignment = 1  # Center alignment

    def _format_heading2(self, paragraph):
        """Format as level 2 heading: HeiBei 14pt bold left."""
        for run in paragraph.runs:
            run.font.name = '黑体'
            run.font.east_asia_font_name = '黑体'
            run.font.size = self.heading2_font_size
            run.bold = True
        paragraph.alignment = 0  # Left alignment

    def _format_heading3(self, paragraph):
        """Format as level 3 heading: Song 12pt bold left."""
        for run in paragraph.runs:
            run.font.name = '宋体'
            run.font.east_asia_font_name = '宋体'
            run.font.size = self.heading3_font_size
            run.bold = True
        paragraph.alignment = 0  # Left alignment

    def _apply_margins(self, doc):
        """Apply GB/T 7714 standard margins to all sections."""
        for section in doc.sections:
            section.left_margin = self.margin_left
            section.right_margin = self.margin_right
            section.top_margin = self.margin_top
            section.bottom_margin = self.margin_bottom

    def _add_page_numbers(self, doc):
        """Add page numbers to footer center of all sections."""
        for section in doc.sections:
            footer = section.footer
            # Clear existing footer content
            for para in footer.paragraphs:
                p = para._element
                p.getparent().remove(p)

            # Create centered paragraph with page number field
            para = footer.add_paragraph()
            para.alignment = 1  # Center alignment

            # Add page number field
            run = para.add_run()
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')

            instrText = OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')
            instrText.text = "PAGE"

            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')

            run._r.append(fldChar1)
            run._r.append(instrText)
            run._r.append(fldChar2)
