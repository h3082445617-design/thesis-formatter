# src/utils/section_detector.py
"""自动检测论文各部分边界"""

from typing import Dict, Tuple, Optional
from docx import Document


class SectionDetector:
    """论文部分检测器"""

    # 部分关键词映射
    KEYWORDS = {
        'abstract': ['摘要', '内容摘要', '中文摘要'],
        'toc': ['目录', '目  录', '目    录'],
        'body': [],  # 通常在摘要和参考文献之间
        'references': ['参考文献', '参考资料', '文献', '参考文献 '],
        'appendix': ['附录', '附  录'],
        'acknowledgment': ['致谢', '后记', '致  谢', '致     谢']
    }

    def __init__(self, doc: Document):
        self.doc = doc
        self.paragraphs = doc.paragraphs

    def detect_parts(self, debug: bool = False) -> Dict[str, Optional[Tuple[int, int]]]:
        """
        检测论文各部分的段落范围

        策略：
        1. 关键词匹配找到每个部分的起始位置
        2. 下一个部分的开始 = 上一个部分的结束
        3. 如果未找到某部分，返回 None

        Returns:
            {
                'abstract': (start_idx, end_idx),
                'toc': (start_idx, end_idx),
                'body': (start_idx, end_idx),
                'references': (start_idx, end_idx),
                'appendix': (start_idx, end_idx),
                'acknowledgment': (start_idx, end_idx)
            }
        """
        parts = {
            'abstract': None,
            'toc': None,
            'body': None,
            'references': None,
            'appendix': None,
            'acknowledgment': None
        }

        # 第一层：关键词匹配
        part_starts = {}  # part_name -> start_idx
        for part_name in parts.keys():
            idx = self._match_keywords(part_name)
            if idx is not None:
                part_starts[part_name] = idx

        if debug:
            print(f"Detected part starts: {part_starts}")

        # 第二层：推断每个部分的范围（start, end）
        # 按逻辑顺序：abstract → toc → body → references → appendix → acknowledgment
        order = ['abstract', 'toc', 'body', 'references', 'appendix', 'acknowledgment']
        part_indices = [(name, part_starts[name]) for name in order if name in part_starts]

        for i, (part_name, start_idx) in enumerate(part_indices):
            # 结束位置 = 下一个部分的开始 - 1（或文档末尾）
            if i + 1 < len(part_indices):
                end_idx = part_indices[i + 1][1] - 1
            else:
                end_idx = len(self.paragraphs) - 1

            parts[part_name] = (start_idx, end_idx)

        return parts

    def _match_keywords(self, part_name: str) -> Optional[int]:
        """匹配关键词，返回第一个匹配的段落索引"""
        keywords = self.KEYWORDS.get(part_name, [])

        for idx, paragraph in enumerate(self.paragraphs):
            text = paragraph.text.strip()
            for keyword in keywords:
                if keyword in text or text == keyword:
                    return idx

        return None

    def _heuristic_detect(self, parts: Dict) -> Dict:
        """启发式检测失败部分"""
        # TODO: 实现
        pass
